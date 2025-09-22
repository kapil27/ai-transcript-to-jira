"""Document parsing service for extracting text from uploaded files."""

import os
import io
import tempfile
from typing import Optional, Dict, Any, BinaryIO
from pathlib import Path

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False

from ..exceptions import ValidationError
from ..utils import LoggerMixin


class DocumentParsingService(LoggerMixin):
    """Service for parsing documents and extracting text content."""
    
    def __init__(self):
        """Initialize document parsing service."""
        self.supported_formats = {
            'text/plain': {'extension': '.txt', 'parser': self._parse_text},
            'application/pdf': {'extension': '.pdf', 'parser': self._parse_pdf},
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': {
                'extension': '.docx', 'parser': self._parse_docx
            },
            'application/msword': {'extension': '.doc', 'parser': self._parse_docx}  # Fallback for .doc
        }
        
        # File size limits (in bytes)
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        self.max_text_length = 100000  # 100k characters
        
        # Check available parsers
        self._check_parser_availability()
    
    def _check_parser_availability(self):
        """Check which document parsers are available."""
        if not PDF_AVAILABLE:
            self.logger.warning("PyPDF2 not available - PDF parsing disabled")
            # Remove PDF support if library not available
            self.supported_formats.pop('application/pdf', None)
        
        if not DOCX_AVAILABLE:
            self.logger.warning("python-docx not available - DOCX parsing disabled")
            # Remove DOCX support if library not available
            self.supported_formats.pop('application/vnd.openxmlformats-officedocument.wordprocessingml.document', None)
            self.supported_formats.pop('application/msword', None)
        
        self.logger.info(f"Document parsing available for: {list(self.supported_formats.keys())}")
    
    def get_supported_formats(self) -> Dict[str, Dict[str, str]]:
        """Get list of supported document formats."""
        return {
            mime_type: {
                'extension': info['extension'],
                'description': self._get_format_description(mime_type)
            }
            for mime_type, info in self.supported_formats.items()
        }
    
    def _get_format_description(self, mime_type: str) -> str:
        """Get human-readable description for mime type."""
        descriptions = {
            'text/plain': 'Plain text files (.txt)',
            'application/pdf': 'PDF documents (.pdf)',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'Word documents (.docx)',
            'application/msword': 'Word documents (.doc)'
        }
        return descriptions.get(mime_type, 'Unknown format')
    
    def detect_file_type(self, file_data: bytes, filename: str = None) -> Optional[str]:
        """Detect file type from content and filename."""
        # Try using python-magic for content-based detection
        if MAGIC_AVAILABLE:
            try:
                mime_type = magic.from_buffer(file_data, mime=True)
                if mime_type in self.supported_formats:
                    return mime_type
            except Exception as e:
                self.logger.warning(f"Magic detection failed: {e}")
        
        # Fallback to extension-based detection
        if filename:
            extension = Path(filename).suffix.lower()
            for mime_type, info in self.supported_formats.items():
                if info['extension'] == extension:
                    return mime_type
        
        # Try to detect based on content patterns
        try:
            # Check for PDF signature
            if file_data.startswith(b'%PDF'):
                return 'application/pdf'
            
            # Check for DOCX (ZIP-based) signature
            if file_data.startswith(b'PK') and b'word/' in file_data[:1024]:
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
            # Check if it's plain text (basic heuristic)
            try:
                file_data.decode('utf-8')
                return 'text/plain'
            except UnicodeDecodeError:
                pass
                
        except Exception as e:
            self.logger.warning(f"Content-based detection failed: {e}")
        
        return None
    
    def validate_file(self, file_data: bytes, filename: str = None) -> Dict[str, Any]:
        """Validate uploaded file."""
        validation_result = {
            'is_valid': True,
            'errors': [],
            'warnings': [],
            'file_info': {}
        }
        
        # Check file size
        file_size = len(file_data)
        if file_size > self.max_file_size:
            validation_result['is_valid'] = False
            validation_result['errors'].append(
                f"File size ({file_size:,} bytes) exceeds maximum allowed size ({self.max_file_size:,} bytes)"
            )
        
        # Check if file is empty
        if file_size == 0:
            validation_result['is_valid'] = False
            validation_result['errors'].append("File is empty")
        
        # Detect file type
        detected_type = self.detect_file_type(file_data, filename)
        if not detected_type:
            validation_result['is_valid'] = False
            validation_result['errors'].append("Unsupported file format")
        else:
            validation_result['file_info']['detected_type'] = detected_type
            validation_result['file_info']['description'] = self._get_format_description(detected_type)
        
        validation_result['file_info']['size_bytes'] = file_size
        validation_result['file_info']['filename'] = filename
        
        return validation_result
    
    def parse_document(self, file_data: bytes, filename: str = None) -> Dict[str, Any]:
        """
        Parse document and extract text content.
        
        Args:
            file_data: Binary file content
            filename: Original filename (optional, helps with type detection)
            
        Returns:
            Dictionary containing extracted text and metadata
            
        Raises:
            ValidationError: If file is invalid or parsing fails
        """
        # Validate file first
        validation_result = self.validate_file(file_data, filename)
        if not validation_result['is_valid']:
            error_msg = '; '.join(validation_result['errors'])
            raise ValidationError(f"File validation failed: {error_msg}")
        
        file_info = validation_result['file_info']
        mime_type = file_info['detected_type']
        
        self.logger.info(f"Parsing document: {filename} ({mime_type}, {file_info['size_bytes']:,} bytes)")
        
        try:
            # Get appropriate parser
            parser = self.supported_formats[mime_type]['parser']
            
            # Parse document
            extracted_text = parser(file_data)
            
            # Validate extracted text
            if not extracted_text.strip():
                raise ValidationError("No text content found in document")
            
            # Check text length
            if len(extracted_text) > self.max_text_length:
                self.logger.warning(f"Extracted text is very long ({len(extracted_text):,} chars), truncating...")
                extracted_text = extracted_text[:self.max_text_length] + "\n\n[Content truncated due to length...]"
            
            result = {
                'text': extracted_text,
                'metadata': {
                    'filename': filename,
                    'mime_type': mime_type,
                    'file_size_bytes': file_info['size_bytes'],
                    'text_length': len(extracted_text),
                    'parser_used': parser.__name__
                },
                'success': True
            }
            
            self.logger.info(f"Document parsed successfully: {len(extracted_text):,} characters extracted")
            return result
            
        except Exception as e:
            self.logger.error(f"Document parsing failed: {e}")
            raise ValidationError(f"Failed to parse document: {str(e)}")
    
    def _parse_text(self, file_data: bytes) -> str:
        """Parse plain text file."""
        try:
            # Try UTF-8 first
            return file_data.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Try latin-1 as fallback
                return file_data.decode('latin-1')
            except UnicodeDecodeError:
                # Last resort: ignore errors
                return file_data.decode('utf-8', errors='ignore')
    
    def _parse_pdf(self, file_data: bytes) -> str:
        """Parse PDF file and extract text."""
        if not PDF_AVAILABLE:
            raise ValidationError("PDF parsing not available (PyPDF2 not installed)")
        
        try:
            # Create file-like object from bytes
            pdf_stream = io.BytesIO(file_data)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            # Extract text from all pages
            extracted_text = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        extracted_text.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    self.logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            if not extracted_text:
                raise ValidationError("No text could be extracted from PDF")
            
            return '\n\n'.join(extracted_text)
            
        except Exception as e:
            if isinstance(e, ValidationError):
                raise
            raise ValidationError(f"PDF parsing error: {str(e)}")
    
    def _parse_docx(self, file_data: bytes) -> str:
        """Parse DOCX file and extract text."""
        if not DOCX_AVAILABLE:
            raise ValidationError("DOCX parsing not available (python-docx not installed)")
        
        try:
            # Create temporary file for docx parsing (required by python-docx)
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_data)
                temp_file_path = temp_file.name
            
            try:
                # Parse DOCX document
                doc = Document(temp_file_path)
                
                # Extract text from paragraphs
                extracted_text = []
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        extracted_text.append(text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            extracted_text.append(' | '.join(row_text))
                
                if not extracted_text:
                    raise ValidationError("No text content found in DOCX document")
                
                return '\n\n'.join(extracted_text)
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except OSError:
                    pass
                    
        except Exception as e:
            if isinstance(e, ValidationError):
                raise
            raise ValidationError(f"DOCX parsing error: {str(e)}")
    
    def get_parsing_stats(self) -> Dict[str, Any]:
        """Get document parsing capabilities and statistics."""
        return {
            'supported_formats': self.get_supported_formats(),
            'limits': {
                'max_file_size_bytes': self.max_file_size,
                'max_file_size_mb': round(self.max_file_size / (1024 * 1024), 1),
                'max_text_length': self.max_text_length
            },
            'available_parsers': {
                'pdf': PDF_AVAILABLE,
                'docx': DOCX_AVAILABLE,
                'magic_detection': MAGIC_AVAILABLE
            }
        }
